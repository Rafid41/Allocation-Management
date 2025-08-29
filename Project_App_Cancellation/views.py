from django.shortcuts import render, redirect
from django.http import HttpResponse
from django.contrib import messages
from .models import Allocation_Number, Final_Allocation, Item
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import parse_xml
import io
from django.utils import timezone
from Project_App_History.models import Project_History as History
from django.db.models import Q

def set_paragraph_format(paragraph, space_before=0, space_after=12, alignment=WD_PARAGRAPH_ALIGNMENT.CENTER):
    """Helper function to format paragraphs with spacing, alignment, and font."""
    paragraph.alignment = alignment
    p_format = paragraph.paragraph_format
    p_format.space_before = Pt(space_before)
    p_format.space_after = Pt(space_after)
    p_format.line_spacing = 1.5  # Set line spacing to 1.5
    for run in paragraph.runs:
        run.font.name = "Times New Roman"


def set_table_borders(table):
    """Apply borders to a table in a Word document."""
    tbl = table._tbl  # Get the XML representation of the table

    # Check if the table already has properties, otherwise create them
    tblPr = tbl.find("./w:tblPr", namespaces=tbl.nsmap)
    if tblPr is None:
        tblPr = parse_xml('<w:tblPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
        tbl.insert(0, tblPr)

    # Define borders
    borders = parse_xml(
        r'<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        r'<w:top w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
        r'<w:left w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
        r'<w:bottom w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
        r'<w:right w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
        r'<w:insideH w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
        r'<w:insideV w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
        r'</w:tblBorders>'
    )

    # Remove old borders (if any) and add new borders
    existing_borders = tblPr.find("./w:tblBorders", namespaces=tbl.nsmap)
    if existing_borders is not None:
        tblPr.remove(existing_borders)

    tblPr.append(borders)


def cancellation_view(request):
    allocated_allocations = Allocation_Number.objects.filter(
        Q(status="Allocated") | Q(status="Modified"),
        final_allocation__isnull=False
    ).distinct()
    cancelled_allocations = Allocation_Number.objects.filter(status="Cancelled")

    if request.method == "POST":
        action = request.POST.get("action")

        if action == "cancel_allocation":
            dhaka_time = timezone.localtime(timezone.now())
            allocation_id = request.POST.get("allocated_allocation")
            if allocation_id:
                allocation = Allocation_Number.objects.get(id=allocation_id)
                final_allocations = Final_Allocation.objects.filter(allocation_no=allocation)

                History.objects.filter(allocation_no=allocation.allocation_no).delete()

                # Restore item quantities
                for entry in final_allocations:
                    item = Item.objects.get(id=entry.item_primary_key)
                    item.quantity_of_item += entry.quantity
                    item.save()


                    History.objects.create(
                            allocation_no=allocation.allocation_no,  # Ensure it's an integer
                            pbs=entry.pbs,
                            project=entry.project.projectId,  # Assuming packageId holds the value
                            item=entry.item,
                            unit_of_item=entry.unit_of_item,
                            warehouse=entry.warehouse,
                            quantity=entry.quantity,
                            # price=entry.price,
                            created_at=dhaka_time,
                            status="Cancelled",
                            remarks="Cancelled at: <b>" + dhaka_time.strftime("%Y-%m-%d %I:%M %p") + "</b>",  
                        )
                

                allocation.status = "Cancelled"
                allocation.save()
                messages.success(request, f"Allocation {allocation.allocation_no} has been cancelled.")
                return redirect("Project_App_Cancellation:cancellation")

        elif action == "download_cancelled":
            allocation_id = request.POST.get("cancelled_allocation")
            if allocation_id:
                allocation = Allocation_Number.objects.get(id=allocation_id)
                final_allocations = Final_Allocation.objects.filter(allocation_no=allocation).order_by("project")

                # Create a Word document
                doc = Document()

                # Set 1-inch margins
                sections = doc.sections
                for section in sections:
                    section.top_margin = Pt(72)
                    section.bottom_margin = Pt(72)
                    section.left_margin = Pt(72)
                    section.right_margin = Pt(72)

                # Add centered heading
                heading = doc.add_paragraph()
                heading_run = heading.add_run(f'Cancelled Project Allocation: Allocation Number - {allocation.allocation_no}')
                heading_run.bold = True
                heading_run.font.size = Pt(16)
                heading_run.font.name = "Times New Roman"
                set_paragraph_format(heading, space_before=0, space_after=12)

                # Create table with borders
                table = doc.add_table(rows=1, cols=5)
                set_table_borders(table)
                table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Center the table

                # Add table headers
                hdr_cells = table.rows[0].cells
                headers = ["Project", "Item", "Warehouse", "Quantity", "Date"]
                for i, header in enumerate(headers):
                    hdr_cells[i].text = header
                    hdr_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    hdr_cells[i].paragraphs[0].runs[0].bold = True
                    hdr_cells[i].paragraphs[0].runs[0].font.name = "Times New Roman"

                # Populate table with data
                for entry in final_allocations:
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(entry.project)
                    row_cells[1].text = str(entry.item)
                    row_cells[2].text = str(entry.warehouse)
                    row_cells[3].text = str(entry.quantity)
                    row_cells[4].text = str(entry.created_at.date())

                    for cell in row_cells:
                        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        cell.paragraphs[0].runs[0].font.name = "Times New Roman"

                # Generate response
                response = HttpResponse(content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                response["Content-Disposition"] = f'attachment; filename="Cancelled_Project_Allocation_{allocation.allocation_no}.docx"'
                doc_io = io.BytesIO()
                doc.save(doc_io)
                doc_io.seek(0)
                response.write(doc_io.read())
                return response

    return render(request, "Project_Templates/Project_App_Cancellation/cancellation.html", {
        "allocated_allocations": allocated_allocations,
        "cancelled_allocations": cancelled_allocations,
    })
