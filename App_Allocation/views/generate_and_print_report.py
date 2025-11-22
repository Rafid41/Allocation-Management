from django.shortcuts import render, redirect, get_object_or_404
from django.http import HttpResponse
from django.contrib import messages
from django.utils import timezone
from django.db.models import Q

from App_Allocation.models import Allocation_Number, Final_Allocation

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from docx.oxml.shared import OxmlElement
from PIL import ImageFont, ImageDraw, Image
from docx.oxml import parse_xml

import datetime
import pytz



def set_font(run, font_name):
    run.font.name = font_name
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    # This makes sure complex scripts (like Bengali) use the font too
    rFonts.set(qn("w:cs"), font_name)


def format_quantity(quantity):
    if quantity == int(quantity):
        return str(int(quantity))
    else:
        return str(quantity)


# Mapping for English to Bengali digits
def convert_to_bengali_digits(number):
    en_to_bn_digits = str.maketrans("0123456789", "০১২৩৪৫৬৭৮৯")
    return str(number).translate(en_to_bn_digits)


def individual_allocation_download(request):
    if request.method == "POST":
        allocation_id = request.POST.get("allocation_no")
        if allocation_id:
            allocation = get_object_or_404(Allocation_Number, id=allocation_id)
            entries = Final_Allocation.objects.filter(allocation_no=allocation)
            print(len(entries))  # Debugging line to check the query

            if not entries.exists():
                messages.error(request, "No allocation entries found for this number.")
                return redirect("App_Allocation:individual_allocation_download")

            doc = Document()

            # ✅ Set 1-inch page margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(0.5)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
            # ✅ Set page size to A4# 🔧 Reduce distance from edge to header/footer content
            section.header_distance = Inches(0.2)
            section.footer_distance = Inches(0.2)

            # Header lines
            header_lines = [
                "বাংলাদেশ পল্লী বিদ্যুতায়ন বোর্ড",
                "এমপিএসএস পরিদপ্তর",
                "সদর দপ্তর ভবন, নিকুঞ্জ-২, ঢাকা।",
            ]
            for line in header_lines:
                p = doc.add_paragraph()
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.space_before = Pt(0)
                run = p.add_run(line)
                run.font.size = Pt(12)
                set_font(run, "Nikosh")  # Set Bengali font for header
                # run.bold = True

            header_email = [
                "E-mail: mpssbreb@gmail.com",
            ]
            

            for line in header_email:
                p = doc.add_paragraph()
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.space_before = Pt(0)
                run = p.add_run(line)
                run.font.size = Pt(14)
                set_font(run, "Times New Roman")  # Set eNGLISH font for email
            p.paragraph_format.space_after = Pt(6)

            # Get current year and convert to Bengali last two digits
            current_year = datetime.datetime.now().year
            last_two_digits_bengali = convert_to_bengali_digits(str(current_year)[-2:])

            # Reference number and current date (in Bengali) on the same line
            ref_date_para = doc.add_paragraph()
            ref_date_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

            # Set a right-aligned tab stop (6.5 inches ~ A4 width)
            ref_date_para.paragraph_format.tab_stops.add_tab_stop(
                Pt(468), WD_PARAGRAPH_ALIGNMENT.RIGHT
            )

            # Format current date in Bengali
            # Set the Dhaka timezone
            dhaka_tz = pytz.timezone("Asia/Dhaka")

            # Get current time in Dhaka timezone
            today = timezone.now().astimezone(dhaka_tz).strftime("%d/%m/%Y")
            bengali_date = convert_to_bengali_digits(today)

            # Add text: Reference left, Date right
            ref_text = f"স্মারক নং- ২৭.১২.০০০০.০১৬.৪০.০৩৫.{last_two_digits_bengali}.\tতারিখঃ {bengali_date} খ্রিঃ।"
            run = ref_date_para.add_run(ref_text)
            run.font.size = Pt(12)
            set_font(run, "Nikosh")

            # doc.add_paragraph()

            # Allocation No



            # --- Text setup ---
            alloc_text = f"Allocation No: {allocation.allocation_no}"
            font_size_pt = 12
            dpi = 96
            font_px = int(font_size_pt * dpi / 72)

            # Load font (update path as needed for your server)
            try:
                ttf_path = "C:\\Windows\\Fonts\\times.ttf"
                pil_font = ImageFont.truetype(ttf_path, font_px)
            except Exception:
                pil_font = ImageFont.load_default()

            # Measure text using Pillow (new Pillow >=10 syntax)
            img = Image.new("RGB", (2000, 2000), (255, 255, 255))
            draw = ImageDraw.Draw(img)
            bbox = draw.textbbox((0, 0), alloc_text, font=pil_font)
            text_width_px = bbox[2] - bbox[0]
            text_height_px = bbox[3] - bbox[1]

            # Convert to inches → points (Word uses points)
            inches_w = text_width_px / dpi
            inches_h = text_height_px / dpi
            pad_x_in = 0.12
            pad_y_in = 0.06
            width_pt = (inches_w + 2 * pad_x_in) * 72
            height_pt = (inches_h + 2 * pad_y_in) * 72

            # --- Create paragraph (keep right aligned for fallback) ---
            alloc_para = doc.add_paragraph()
            alloc_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            run = alloc_para.add_run()

            # --- Create right-aligned tight rectangle ---
            textbox_xml = f"""
            <w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                xmlns:v="urn:schemas-microsoft-com:vml"
                xmlns:w10="urn:schemas-microsoft-com:office:word">
            <w:pict>
                <v:shape id="AllocationBox" type="#_x0000_t202" strokecolor="black" strokeweight="1pt" fillcolor="white"
                        style="width:{width_pt}pt;height:{height_pt}pt;position:absolute;right:0pt;margin-right:0pt;white-space:nowrap;mso-position-horizontal:right;">
                <v:textbox inset="2pt,2pt,2pt,2pt">
                    <w:txbxContent>
                    <w:p>
                        <w:r>
                        <w:t xml:space="preserve">{alloc_text}</w:t>
                        </w:r>
                    </w:p>
                    </w:txbxContent>
                </v:textbox>
                </v:shape>
            </w:pict>
            </w:r>
            """

            # Append the shape XML to the run
            run._r.append(parse_xml(textbox_xml))

            # Optional spacing
            alloc_para.paragraph_format.space_after = Pt(6)


            # doc.add_paragraph()

            # Recipient block
            recipients = [
                "পরিচালক",
                "সিএসএন্ডএম পরিদপ্তর,",
                "বাপবিবো, ঢাকা।",
            ]
            for recipient in recipients:
                p = doc.add_paragraph(recipient)
                p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.space_before = Pt(0)
                for run in p.runs:
                    run.font.size = Pt(12)
                    set_font(run, "Nikosh")
                    # run.bold = True
            p.paragraph_format.space_after = Pt(6)

            # doc.add_paragraph()

            subj = doc.add_paragraph()
            run = subj.add_run("    বিষয় :  মালামাল বরাদ্দকরণ প্রসংগে।")
            run.bold = True
            run.font.size = Pt(12)
            set_font(run, "Nikosh")

            # doc.add_paragraph()

            # Parts of the paragraph before, the English phrase, and after
            part1 = "বাপবিবোর্ডের “"
            english_text = "PBS Fund (084)"
            part2 = (
                "” এর আওতায় সংগ্রহকৃত মালামাল নিম্নের ছকে উল্লেখিত কেন্দ্রীয় পন্যাগার হতে মূল্য পরিশোধ স্বাপেক্ষে "
                "নিম্নোক্তভাবে বরাদ্দ প্রদান করা হলো। বরাদ্দপ্রাপ্ত মালামাল এর মূল্য বাপবিবো’র হিসাব পরিদপ্তরে জমা প্রদানের রশিদ দাখিল পূর্বক "
                "নিম্নের ছকে উল্লেখিত কেন্দ্রীয় পন্যাগার হতে মালামালসমূহ গ্রহন করবে। "
            )

            bold_part = "তবে মূল্য পরিশোধের পূর্বে কেন্দ্রীয় পণ্যাগারে আইটেমটির মজুদের তথ্য নেওয়ার জন্য অনুরোধ করা হলো।"

            para = doc.add_paragraph()

            # Run 1: Bengali before English phrase
            run1 = para.add_run(part1)
            run1.font.size = Pt(12)
            set_font(run1, "Nikosh")

            # Run 2: English phrase in Times New Roman
            run2 = para.add_run(english_text)
            run2.font.size = Pt(12)
            run2.font.name = "Times New Roman"
            # For Word to recognize font properly (sometimes needed)
            rFonts = run2._element.rPr.rFonts
            rFonts.set(qn("w:eastAsia"), "Times New Roman")

            # Run 3: Bengali after English phrase
            run3 = para.add_run(part2)
            run3.font.size = Pt(12)
            set_font(run3, "Nikosh")

            # Run 4: Bold Bengali part
            run4 = para.add_run(bold_part)
            run4.bold = True
            run4.font.size = Pt(12)
            set_font(run4, "Nikosh")

            para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

            # doc.add_paragraph()

            # Create table
            table = doc.add_table(rows=1, cols=7)
            table.style = "Table Grid"

            # Set column headers
            hdr_cells = table.rows[0].cells
            headers = [
                "গ্রহণকারী সমিতির নাম",
                "আইটেম নং",
                "পরিমাণ",
                "একক",
                "একক মূল্য\n(টাকা)",
                "বরাদ্দকৃত প্রকল্প ও প্যাকেজ/ সাব-প্যাকেজ নং",
                "পণ্যাগারের নাম",
            ]

            for i, text in enumerate(headers):
                hdr_cells[i].text = text
                hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                para = hdr_cells[i].paragraphs[0]
                para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER 

                for run in para.runs:
                    run.font.size = Pt(12)
                    run.bold = True
                    run.font.name = "Nikosh"
                    # Ensure Bengali font for complex scripts
                    rFonts = run._element.rPr.get_or_add_rFonts()
                    rFonts.set(qn("w:cs"), "Nikosh")


            # Widen column 5
            hdr_cells[4].width = Inches(2.5)

            # first_col8_cell = None

            for idx, entry in enumerate(entries):
                row_cells = table.add_row().cells

                row_cells[0].width = Inches(3.0)
                row_cells[1].width = Inches(4.0)
                row_cells[2].width = Inches(1.5) # Quantity
                row_cells[3].width = Inches(1.5) # Unit
                row_cells[4].width = Inches(3.0) # Price
                row_cells[5].width = Inches(7.0) # Package
                row_cells[6].width = Inches(2.5) # Warehouse

                values = [
                    str(entry.pbs.name),
                    str(entry.item),
                    format_quantity(entry.quantity),
                    str(entry.unit_of_item),
                    str(entry.price),
                    str(entry.package.packageId),
                    "CWH, " + str(entry.warehouse)
                ]

                for i, value in enumerate(values):
                    row_cells[i].text = value
                    row_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    para = row_cells[i].paragraphs[0]
                    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    for run in para.runs:
                        run.font.size = Pt(11)
                        run.font.name = "Times New Roman"
                        rFonts = run._element.rPr.get_or_add_rFonts()
                        rFonts.set(qn("w:cs"), "Times New Roman")

                # if idx == 0:
                #     first_col8_cell = row_cells[7]
                # else:
                #     first_col8_cell.merge(row_cells[7])


            # Set value and formatting for the merged column 8 cell
            # first_col8_cell.text = "On Payment O&M Store"
            # first_col8_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            # para = first_col8_cell.paragraphs[0]
            # para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # # Ensure the paragraph has text and formatting
            # if not para.runs:
            #     run = para.add_run("On Payment O&M Store")
            # else:
            #     para.runs[0].text = "On Payment O&M Store"
            #     run = para.runs[0]

            # run.font.size = Pt(11)
            # run.font.name = "Times New Roman"
            # rFonts = run._element.rPr.get_or_add_rFonts()
            # rFonts.set(qn("w:cs"), "Times New Roman")


            # doc.add_paragraph()

            closing_text = "বরাদ্দ অনুযায়ী প্রয়োজনীয় ব্যবস্থা গ্রহণের অনুরোধ করা হলো।"
            closing_para = doc.add_paragraph(closing_text)
            closing_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            for run in closing_para.runs:
                run.font.size = Pt(12)
                set_font(run, "Nikosh")

            # doc.add_paragraph()

            sigs = [
                # Right-aligned block
                {"text": "(       )     ", "align": WD_PARAGRAPH_ALIGNMENT.RIGHT},
                {"text": "উপ-পরিচালক(কারিগরী)", "align": WD_PARAGRAPH_ALIGNMENT.RIGHT},
                {"text": "", "align": WD_PARAGRAPH_ALIGNMENT.RIGHT},
                # Left-aligned block
                {"text": "অনুলিপি:", "align": WD_PARAGRAPH_ALIGNMENT.LEFT},
                {
                    "text": "১।      উপ-পরিচালক, কেন্দ্রীয় পণ্যাগার, বাপবিবো,.......।",
                    "align": WD_PARAGRAPH_ALIGNMENT.LEFT,
                },
                {
                    "text": "২।      উপ-পরিচালক, মালামাল হিসাব, বাপবিবো, ঢাকা।",
                    "align": WD_PARAGRAPH_ALIGNMENT.LEFT,
                },
                {
                    "text": "৩।      জেনারেল ম্যানেজার,..... পবিস।",
                    "align": WD_PARAGRAPH_ALIGNMENT.LEFT,
                },
                # Right-aligned block
                {"text": "", "align": WD_PARAGRAPH_ALIGNMENT.RIGHT},
                {"text": "(          )     ", "align": WD_PARAGRAPH_ALIGNMENT.RIGHT},
                {"text": "সহকারী প্রকৌশলী", "align": WD_PARAGRAPH_ALIGNMENT.RIGHT},
            ]

            for sig in sigs:
                p = doc.add_paragraph(sig["text"])
                p.alignment = sig["align"]

                # 🔧 Remove spacing after/before paragraph
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.space_before = Pt(0)

                for run in p.runs:
                    set_font(run, font_name="Nikosh")
                    if sig["text"] == "অনুলিপি:":
                        run.underline = True

            response = HttpResponse(
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            filename = f"Allocation_Report_{allocation.allocation_no}.docx"
            response["Content-Disposition"] = f'attachment; filename="{filename}"'
            doc.save(response)
            return response

        messages.error(request, "Please select an allocation number.")
        return redirect("App_Allocation:individual_allocation_download")

    # allocated_allocations = Allocation_Number.objects.filter(
    #     Q(status="Allocated") | Q(status="Modified"),
    #     allocation_no__in=Final_Allocation.objects.values_list(
    #         "allocation_no", flat=True
    #     ).distinct(),
    # ).order_by("-allocation_no")
    allocated_allocations = Allocation_Number.objects.filter(
        Q(status="Allocated") | Q(status="Modified"),
        id__in=Final_Allocation.objects.values_list("allocation_no_id", flat=True).distinct(),
    ).order_by("-allocation_no")

    return render(
        request,
        "App_Allocation/generate_and_print_report.html",
        {"allocated_allocations": allocated_allocations},
    )