from django.shortcuts import render, redirect
from django.db.models import Q
from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger
from django.http import JsonResponse, HttpResponse
from django.contrib.auth.decorators import login_required
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from io import BytesIO
from .models import PBS_History
from App_Admin_Panel.models import PaginationManager

@login_required
def pbswise_history_view(request):
    """
    Authorized view for tracking longitudinal PBS inventory movements.
    Restricted to Superusers and Editors for administrative auditing.
    """
    if not request.user.is_superuser and request.user.user_group.user_group_type != "Editor":
        return redirect("App_User_Group:access-denied")

    # Initial Queryset
    queryset = PBS_History.objects.select_related('pbs', 'item', 'zonal_from', 'zonal_to').all()

    # --- Filtering Logic ---
    pbs_name = request.GET.get('pbs_name', '').strip()
    item_name = request.GET.get('item_name', '').strip()
    action = request.GET.get('action', '').strip()
    zonal_from = request.GET.get('zonal_from', '').strip()
    store_from = request.GET.get('store_from', '').strip()
    date_val = request.GET.get('date', '').strip()

    if pbs_name:
        queryset = queryset.filter(pbs__pbs_name=pbs_name)
    if item_name:
        queryset = queryset.filter(item__item_name=item_name)
    if action and action != "Any action":
        queryset = queryset.filter(action=action)
    if zonal_from:
        queryset = queryset.filter(zonal_from__zonal_name=zonal_from)
    if store_from:
        queryset = queryset.filter(store_from=store_from)
    if date_val:
        # Datepicker format YYYY-MM-DD
        queryset = queryset.filter(date__date=date_val)

    # Dynamic Pagination / Print View Controller
    is_print_view = request.GET.get('print_view') == 'true'
    
    if is_print_view:
        items = queryset
    else:
        try:
            limit = PaginationManager.load().table_pagination_limit
        except:
            limit = 50

        paginator = Paginator(queryset, limit)
        page_number = request.GET.get('page', 1)
        try:
            items = paginator.page(page_number)
        except PageNotAnInteger:
            items = paginator.page(1)
        except EmptyPage:
            items = paginator.page(paginator.num_pages)

    context = {
        'items': items,
        'pbs_name': pbs_name,
        'item_name': item_name,
        'selected_action': action,
        'zonal_from_val': zonal_from,
        'store_from_val': store_from,
        'date_val': date_val,
        'is_print_view': is_print_view,
        'actions': ['Transfer Item', 'Withdraw Item']
    }
    return render(request, "PBSWise_Templates/PBSWise_History/PBSWise_History.html", context)

def get_history_suggestions_ajax(request):
    """AJAX helper for fetching absolute unique suggestions from available history data."""
    field = request.GET.get('field')
    query = request.GET.get('query', '')
    
    pbs_name = request.GET.get('pbs_name', '')
    zonal_from = request.GET.get('zonal_from', '')

    queryset = PBS_History.objects.all()
    suggestions = []
    
    if field == 'pbs_name':
        suggestions = queryset.filter(pbs__pbs_name__icontains=query).values_list('pbs__pbs_name', flat=True).order_by('pbs__pbs_name').distinct()
    elif field == 'item_name':
        suggestions = queryset.filter(item__item_name__icontains=query).values_list('item__item_name', flat=True).order_by('item__item_name').distinct()
    elif field == 'zonal_from':
        if pbs_name:
            queryset = queryset.filter(pbs__pbs_name=pbs_name)
        suggestions = queryset.filter(zonal_from__zonal_name__icontains=query).values_list('zonal_from__zonal_name', flat=True).order_by('zonal_from__zonal_name').distinct()
    elif field == 'store_from':
        if pbs_name:
            queryset = queryset.filter(pbs__pbs_name=pbs_name)
        if zonal_from:
            queryset = queryset.filter(zonal_from__zonal_name=zonal_from)
        suggestions = queryset.filter(store_from__icontains=query).values_list('store_from', flat=True).order_by('store_from').distinct()

    return JsonResponse({'suggestions': list(suggestions)[:10]})

@login_required
def export_history_docx(request):
    """Generates a professional administrative Word document with regional project formatting."""
    document = Document()
    
    # 1. Page Margins (0.75 inches for regional project standards)
    section = document.sections[0]
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)

    # 2. Font Setup (Times New Roman)
    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)

    # 3. Bold, Centered Headline
    headline = document.add_heading('PBS Inventory Operation History', 0)
    headline.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Apply same filters as main view if passed
    pbs_name = request.GET.get('pbs_name', '')
    item_name = request.GET.get('item_name', '')
    
    queryset = PBS_History.objects.select_related('pbs', 'item', 'zonal_from', 'zonal_to').all()
    if pbs_name: queryset = queryset.filter(pbs__pbs_name=pbs_name)
    if item_name: queryset = queryset.filter(item__item_name=item_name)

    # 4. Professional Table Suite
    table = document.add_table(rows=1, cols=9)
    table.style = 'Table Grid'
    table.autofit = False
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Proportional Cell Widths (Total 7.0 inches for 8.5" page with 0.75*2 margins)
    col_widths = [1.1, 0.7, 1.2, 0.4, 0.7, 0.7, 0.7, 0.7, 0.7]
    
    hdr_cells = table.rows[0].cells
    headers = ['PBS Name', 'Date', 'Item', 'Qty', 'Action', 'Zonal From', 'Store From', 'Zonal To', 'Store To']
    
    for i, header_text in enumerate(headers):
        cell = hdr_cells[i]
        cell.width = Inches(col_widths[i])
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(header_text)
        run.bold = True
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for rec in queryset:
        row_cells = table.add_row().cells
        data = [
            str(rec.pbs.pbs_name),
            rec.date.strftime('%d-%m-%Y'),
            str(rec.item.item_name),
            str(f"{rec.quantity:g}"),
            str(rec.action),
            str(rec.zonal_from.zonal_name),
            str(rec.store_from).replace('_', ' ').title(),
            str(rec.zonal_to.zonal_name) if rec.zonal_to else 'N/A',
            str(rec.store_to).replace('_', ' ').title() if rec.store_to else 'N/A'
        ]
        for i, val in enumerate(data):
            cell = row_cells[i]
            cell.width = Inches(col_widths[i])
            cell.text = val
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT if i in [0, 2] else WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    f = BytesIO()
    document.save(f)
    f.seek(0)

    response = HttpResponse(f.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=PBS_History_Export.docx'
    return response
