# views.py
from django.shortcuts import render, redirect
from django.http import HttpResponse
from django.contrib import messages
from App_Allocation.models import Allocation_Number, Final_Allocation
from App_History.models import History
from django.db.models import Q
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import parse_xml
import io

def set_paragraph_format(paragraph, space_before=0, space_after=12, alignment=WD_PARAGRAPH_ALIGNMENT.CENTER):
    paragraph.alignment = alignment
    p_format = paragraph.paragraph_format
    p_format.space_before = Pt(space_before)
    p_format.space_after = Pt(space_after)
    p_format.line_spacing = 1.5
    for run in paragraph.runs:
        run.font.name = "Times New Roman"

def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.find("./w:tblPr", namespaces=tbl.nsmap)
    if tblPr is None:
        tblPr = parse_xml('<w:tblPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
        tbl.insert(0, tblPr)

    borders = parse_xml(
        r'<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        r'<w:top w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
        r'<w:left w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
        r'<w:bottom w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
        r'<w:right w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
        r'<w:insideH w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
        r'<w:insideV w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
        r'</w:tblBorders>'
    )
    existing_borders = tblPr.find("./w:tblBorders", namespaces=tbl.nsmap)
    if existing_borders is not None:
        tblPr.remove(existing_borders)
    tblPr.append(borders)

def modification_view(request):
    allocated_allocations = Allocation_Number.objects.filter(
        Q(status="Allocated") | Q(status="Modified"),
        final_allocation__isnull=False
    ).distinct()

    modified_allocations = Allocation_Number.objects.filter(status="Modified")

    if request.method == "POST":
        action = request.POST.get("action")

        if action == "go_to_modify":
            allocation_id = request.POST.get("allocated_allocation")
            if allocation_id:
                return redirect("App_Modification:modification_options", allocation_id=allocation_id)

        elif action == "download_modified":
            allocation_id = request.POST.get("modified_allocation")
            if allocation_id:
                allocation = Allocation_Number.objects.get(id=allocation_id)
                final_allocations = Final_Allocation.objects.filter(allocation_no=allocation).order_by("package")

                doc = Document()

                for section in doc.sections:
                    section.top_margin = Pt(72)
                    section.bottom_margin = Pt(72)
                    section.left_margin = Pt(72)
                    section.right_margin = Pt(72)

                heading = doc.add_paragraph()
                heading_run = heading.add_run(f'Modified Allocation: Allocation Number - {allocation.allocation_no}')
                heading_run.bold = True
                heading_run.font.size = Pt(16)
                heading_run.font.name = "Times New Roman"
                set_paragraph_format(heading, space_before=0, space_after=12)

                table = doc.add_table(rows=1, cols=6)
                set_table_borders(table)
                table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                headers = ["Package", "Item", "Warehouse", "Quantity", "Price", "Date"]
                hdr_cells = table.rows[0].cells
                for i, header in enumerate(headers):
                    hdr_cells[i].text = header
                    hdr_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    hdr_cells[i].paragraphs[0].runs[0].bold = True
                    hdr_cells[i].paragraphs[0].runs[0].font.name = "Times New Roman"

                for entry in final_allocations:
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(entry.package)
                    row_cells[1].text = str(entry.item)
                    row_cells[2].text = str(entry.warehouse)
                    row_cells[3].text = str(entry.quantity)
                    row_cells[4].text = str(entry.price)
                    row_cells[5].text = str(entry.created_at.date())

                    for cell in row_cells:
                        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        cell.paragraphs[0].runs[0].font.name = "Times New Roman"

                response = HttpResponse(content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                response["Content-Disposition"] = f'attachment; filename="Modified_Allocation_{allocation.allocation_no}.docx"'
                doc_io = io.BytesIO()
                doc.save(doc_io)
                doc_io.seek(0)
                response.write(doc_io.read())
                return response

    return render(request, "App_Modification/modification.html", {
        "allocated_allocations": allocated_allocations,
        "modified_allocations": modified_allocations,
    })
    